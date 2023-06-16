from docx import Document
import pandas as pd
from datetime import *

import tkinter as tk

from tkinter import *
from tkinter import ttk

import win32com.client as win32



Documento=Document("Documento.docx")
"EEE" #nome da escola
"END" #endereço da escola
"NR" #nome responsavel
"ER" #endereço responsavel
"CR" #cpf do responsavel


"NA"#nome do aluno
"DA" #data de nascimento do aluno
"GA" #genero do aluno
"EA" #endereco do aluno
"TEL" #telefone de contato

"TA" #turma
"AL" #ano letivo

"PA"#pais do aluno
"CID" #cidade
"PN" #periodo de notificações
for paragrafos in Documento.paragraphs:
    print(paragrafos.text)

def salvar():
    nomeR=str(NREntry.get())
    NREntry.delete("0","end")

    endR=str(EREntry.get())
    EREntry.delete("0","end")

    cpfR=str(CREntry.get())
    CREntry.delete("0","end")

    nomeA=str(NAEntry.get())
    NAEntry.delete("0","end")

    dataA=str(DNAEntry.get())
    DNAEntry.delete("0","end")

    endA=str(EAEntry.get())
    EAEntry.delete("0","end")

    genA=str(GAEntry.get())
    GAEntry.delete("0","end")

    telR=str(TEAEntry.get())
    TEAEntry.delete("0","end")

    turma=str(TAEntry.get())
    TAEntry.delete("0","end")

    cidade=str(CIDEntry.get())
    CIDEntry.delete("0","end")

    dados={
    "EEE":"ESCOLA FICTICIA",
    "END":"na rua ficticia",
    "NR":nomeR,
    "ER":endR,
    "CR":cpfR,
    "NA":nomeA,
    "DA":dataA,
    "EA":endA,
    "TEL":telR,
    "TA": turma+'º',
    "AL": str(datetime.now().year),
    "CID":cidade,
    "PA":"Brasil",
    "GA":genA,
    }

    for paragrafo in Documento.paragraphs:
        for codigo in dados:
            valor=dados[codigo]
            paragrafo.text=paragrafo.text.replace(codigo, valor)

    for paragrafo in Documento.paragraphs:
        print(paragrafo.text)

    Documento.save(f'Contrato-{nomeA}.docx')

    newWindow=Tk()
    newWindow.iconbitmap("python_18894.ico")
    newWindow.maxsize(350,80)
    newWindow.minsize(350,80)
    Text=Label(newWindow,text="Salvo em docx na pasta dessa executavel.")
    Text.grid(column=0,row=0)
        
    newWindow.mainloop()


    

window= Tk()


window.title("Escola-APP")
window.iconbitmap('python_18894.ico')


window.geometry("") 
window.minsize(800, 250) 
window.maxsize(800, 250) 

NR=Label(window,text="Nome do Responsável:")
NR.grid(column=0,row=0)
NREntry=Entry(window)
NREntry.grid(column=1,row=0)

ER=Label(window,text="Endereço do Responsavel:")
ER.grid(column=0,row=1)
EREntry=Entry(window)
EREntry.grid(column=1,row=1)

CR=Label(window,text="CPF do responsavel:")
CR.grid(column=0,row=2)
CREntry=Entry(window)
CREntry.grid(column=1,row=2)

NA=Label(window, text="Nome do Aluno:")
NA.grid(column=5,row=0)
NAEntry=Entry(window)
NAEntry.grid(column=6,row=0)

DNA=Label(window, text="Nascimento:")
DNA.grid(column=5,row=1)
DNAEntry=Entry(window)
DNAEntry.grid(column=6,row=1)

GA=Label(window, text="Genero do Aluno:")
GA.grid(column=5,row=2)
GAEntry=Entry(window)
GAEntry.grid(column=6,row=2)

EA=Label(window, text="Endereço do Aluno:")
EA.grid(column=5,row=3)
EAEntry=Entry(window)
EAEntry.grid(column=6,row=3)

TEA=Label(window, text="Telefone:")
TEA.grid(column=3,row=4)
TEAEntry=Entry(window)
TEAEntry.grid(column=4,row=4)

TA=Label(window, text="Turma:")
TA.grid(column=3,row=5)
TAEntry=Entry(window)
TAEntry.grid(column=4,row=5)

CID=Label(window, text="Cidade:")
CID.grid(column=3,row=6)
CIDEntry=Entry(window)
CIDEntry.grid(column=4,row=6)


botao=Button(window,text="Fazer arquivo", command=salvar)
botao.grid(column=0,row=10)





window.mainloop()
